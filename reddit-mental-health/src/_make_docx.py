"""Convert METHODOLOGY.md to METHODOLOGY.docx using python-docx.

Handles: headings (# ## ###), tables (|---|), bold (**text**),
inline code (`text`), code blocks (```), blockquotes (>), bullet
lists (- / *), numbered lists (1.), horizontal rules (---), and
mixed inline formatting within paragraphs.
"""
from __future__ import annotations
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
MD_IN  = ROOT / "METHODOLOGY.md"
DOCX_OUT = ROOT / "METHODOLOGY.docx"


# ── Helpers ──────────────────────────────────────────────────────────────────

def add_horizontal_rule(doc: Document) -> None:
    """Add a thin horizontal line (paragraph border)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_bg(cell, hex_color: str) -> None:
    """Set table cell background shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def apply_inline(run, text: str) -> None:
    """Set run text, stripping any stray backticks left by the caller."""
    run.text = text


def add_inline_text(para, raw: str) -> None:
    """
    Parse a line of markdown inline markup and add styled runs to `para`.

    Handles: **bold**, *italic*, `code`, and plain text.
    Multiple patterns can appear on the same line.
    """
    # Pattern: captures bold, italic, inline-code, or plain text segments
    token_re = re.compile(
        r"(\*\*(.+?)\*\*)"          # bold
        r"|(\*(.+?)\*)"             # italic
        r"|(`(.+?)`)"               # inline code
        r"|(\$\$.+?\$\$)"           # display math  → render as plain
        r"|(\$.+?\$)"               # inline math   → render as plain
        r"|([^*`$]+)"               # plain text
    )
    for m in token_re.finditer(raw):
        if m.group(1):              # **bold**
            run = para.add_run(m.group(2))
            run.bold = True
        elif m.group(3):            # *italic*
            run = para.add_run(m.group(4))
            run.italic = True
        elif m.group(5):            # `code`
            run = para.add_run(m.group(6))
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif m.group(7):            # $$math$$
            run = para.add_run(m.group(7))
            run.italic = True
        elif m.group(8):            # $math$
            run = para.add_run(m.group(8))
            run.italic = True
        elif m.group(9):            # plain
            para.add_run(m.group(9))


# ── Table parsing ─────────────────────────────────────────────────────────────

def parse_table_block(lines: list[str]) -> list[list[str]]:
    """Parse a markdown table into a list-of-rows list-of-cells."""
    rows: list[list[str]] = []
    for line in lines:
        if re.match(r"^\s*\|?[-:| ]+\|?\s*$", line):
            continue  # separator row
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        row = table.add_row()
        # Header row: bold + light blue bg
        is_header = i == 0
        for j, cell_text in enumerate(row_data):
            if j >= ncols:
                break
            cell = row.cells[j]
            if is_header:
                set_cell_bg(cell, "DCE6F1")
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline_text(para, cell_text)
            if is_header:
                for run in para.runs:
                    run.bold = True
            # Smaller font inside table
            for run in para.runs:
                run.font.size = Pt(9.5)
    doc.add_paragraph()  # spacing after table


# ── Code block ───────────────────────────────────────────────────────────────

def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        para = doc.add_paragraph(style="No Spacing")
        run = para.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x23, 0x23, 0x23)
    doc.add_paragraph()


# ── Main converter ────────────────────────────────────────────────────────────

def convert(md_path: Path, docx_path: Path) -> None:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.15)
        section.right_margin  = Inches(1.15)

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]

        # ── Fenced code block ───────────────────────────────────────────
        if line.strip().startswith("```"):
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines)
            i += 1
            continue

        # ── Horizontal rule (--- alone on a line, not a table sep) ──────
        if re.match(r"^-{3,}\s*$", line) and not line.strip().startswith("|"):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Table block (collect all consecutive table lines) ────────────
        if re.match(r"^\s*\|", line):
            table_lines: list[str] = []
            while i < len(lines) and re.match(r"^\s*\|", lines[i]):
                table_lines.append(lines[i])
                i += 1
            rows = parse_table_block(table_lines)
            add_table(doc, rows)
            continue

        # ── Headings ─────────────────────────────────────────────────────
        h_match = re.match(r"^(#{1,4})\s+(.*)", line)
        if h_match:
            level  = len(h_match.group(1))
            text   = h_match.group(2)
            # Strip trailing ==/ -- setext markers (shouldn't appear but safe)
            text   = re.sub(r"\*\*(.+?)\*\*", r"\1", text)  # strip bold in headings
            style_map = {1: "Heading 1", 2: "Heading 2",
                         3: "Heading 3", 4: "Heading 4"}
            para = doc.add_heading(text, level=level)
            i += 1
            continue

        # ── Blockquote ───────────────────────────────────────────────────
        if line.startswith(">"):
            text = line.lstrip("> ").strip()
            para = doc.add_paragraph(style="Quote")
            add_inline_text(para, text)
            para.paragraph_format.left_indent  = Inches(0.4)
            para.paragraph_format.right_indent = Inches(0.4)
            for run in para.runs:
                run.italic = True
            i += 1
            continue

        # ── Bullet list ──────────────────────────────────────────────────
        if re.match(r"^(\s*)[-*]\s+", line):
            indent = len(re.match(r"^(\s*)", line).group(1))
            text   = re.sub(r"^\s*[-*]\s+", "", line)
            para   = doc.add_paragraph(style="List Bullet")
            add_inline_text(para, text)
            if indent >= 2:
                para.paragraph_format.left_indent = Inches(0.5)
            i += 1
            continue

        # ── Numbered list ────────────────────────────────────────────────
        if re.match(r"^\d+\.\s+", line):
            text = re.sub(r"^\d+\.\s+", "", line)
            para = doc.add_paragraph(style="List Number")
            add_inline_text(para, text)
            i += 1
            continue

        # ── Empty line → paragraph break ────────────────────────────────
        if line.strip() == "":
            i += 1
            continue

        # ── Regular paragraph ────────────────────────────────────────────
        para = doc.add_paragraph()
        add_inline_text(para, line.strip())
        para.paragraph_format.space_after = Pt(6)
        i += 1

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == "__main__":
    convert(MD_IN, DOCX_OUT)
